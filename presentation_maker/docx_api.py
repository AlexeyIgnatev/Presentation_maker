from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from presentation_maker.saver import check_path, path_to_file
from presentation_maker.config import source_path
import os


class Docx:
    def __init__(self):
        self.document = None
        self.topic = ''
        self.file = ''

    def create_document(self, text: str):
        document = Document()
        text = text.split('\n')
        self.topic = text[0].replace('=', '').strip()
        for i in text:
            if i.startswith('='):
                length = list(i.split()[0]).count('=') - 1
                p = document.add_paragraph()
                run = p.add_run(i)
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                font = run.font
                font.name = 'Calibri'
                font.size = Pt(20 - length * 2)
                run.bold = True

            else:
                document.add_paragraph(i)
        self.document = document

    def save_document(self):
        if self.document is None:
            print('Создайте документ перед соханением!')
        else:
            check_path(source_path)
            check_path('{}/{}'.format(source_path, self.topic))
            self.file = '{}/{}/{}.docx'.format(source_path, self.topic, self.topic)
            self.document.save(self.file)

    def open_document(self):
        if self.file:
            os.startfile(path_to_file(self.file))
